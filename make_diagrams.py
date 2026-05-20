"""
Generate all project diagrams as PNG files, then bundle into a Word doc.

Output files:
  diagrams/gantt_chart.png
  diagrams/architecture_diagram.png
  diagrams/data_processing_workflow.png
  diagrams/team_structure.png
  FACEE_Diagrams.docx
"""

import os, math
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import matplotlib.patheffects as pe
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.shared import Inches, Pt, RGBColor as DRGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION

os.makedirs("diagrams", exist_ok=True)

WHITE  = '#FFFFFF'
BLACK  = '#000000'
DBLU   = '#1F497D'
MBLU   = '#2E74B5'
LGRAY  = '#D6DCE4'
GRAY   = '#808080'
RED_G  = '#FF0000'
GREEN_G = '#00B050'
ORANGE = '#ED7D31'


# ═══════════════════════════════════════════════════════════════
# 1. GANTT CHART
# ═══════════════════════════════════════════════════════════════
def make_gantt():
    COL_W = [2.2] + [1.3]*6
    COL_X = [0.0]
    for w in COL_W:
        COL_X.append(COL_X[-1] + w)
    HDR_H, ROW_H = 0.72, 0.47
    ROW_Y = [0.0, HDR_H]
    for _ in range(14):
        ROW_Y.append(ROW_Y[-1] + ROW_H)
    TW, TH = COL_X[-1], ROW_Y[-1]

    fig, ax = plt.subplots(figsize=(12, 7))
    ax.set_xlim(0, TW); ax.set_ylim(0, TH)
    ax.invert_yaxis(); ax.axis('off')
    fig.patch.set_facecolor(WHITE)

    def rect(x, y, w, h, color, z=1):
        ax.add_patch(mpatches.Rectangle((x, y), w, h,
                     facecolor=color, edgecolor='none', zorder=z))

    def fc(col, row, color, fs=0.0, fe=1.0):
        cw = COL_X[col+1]-COL_X[col]
        rect(COL_X[col]+fs*cw, ROW_Y[row], (fe-fs)*cw,
             ROW_Y[row+1]-ROW_Y[row], color)

    def ar(ai, sub): return 1 + ai*2 + sub

    rect(0, 0, TW, TH, WHITE)
    R, G = RED_G, GREEN_G

    fc(1,ar(0,0),R); fc(2,ar(0,0),R); fc(1,ar(0,1),G); fc(2,ar(0,1),G)
    fc(1,ar(1,0),R); fc(2,ar(1,0),R); fc(4,ar(1,0),R); fc(5,ar(1,0),R)
    fc(1,ar(1,1),G); fc(2,ar(1,1),G)
    fc(2,ar(2,0),R); fc(3,ar(2,0),R,fe=0.60); fc(2,ar(2,1),G,fe=0.65)
    fc(2,ar(3,0),R,fs=0.55); fc(3,ar(3,0),R); fc(4,ar(3,0),R); fc(5,ar(3,0),R)
    fc(2,ar(3,1),G,fs=0.55); fc(3,ar(3,1),G,fe=0.28)
    fc(4,ar(4,0),R); fc(5,ar(4,0),R)
    fc(5,ar(5,0),R); fc(6,ar(5,0),R,fe=0.38)
    fc(1,ar(6,0),R); fc(2,ar(6,0),R); fc(4,ar(6,0),R); fc(5,ar(6,0),R); fc(6,ar(6,0),R)
    fc(1,ar(6,1),G); fc(2,ar(6,1),G)

    LW = 1.8
    for cx in COL_X: ax.plot([cx,cx],[0,TH],'-',color=BLACK,lw=LW,zorder=2)
    for ry in ROW_Y: ax.plot([0,TW],[ry,ry],'-',color=BLACK,lw=LW,zorder=2)

    HDRS = ["Activities","28 Jan-\n7 Feb","8 Feb-\n19 Mar","20 Mar-\n2 May",
            "3 May-\n6 Jun","7 Jun-\n4 July","5 Jul-10\nJul"]
    for ci, h in enumerate(HDRS):
        cw = COL_X[ci+1]-COL_X[ci]
        xp = COL_X[ci]+(0.15 if ci==0 else cw/2)
        ax.text(xp, HDR_H/2, h, ha='left' if ci==0 else 'center', va='center',
                fontsize=10, fontweight='bold', color=BLACK, linespacing=1.25, zorder=3)

    ACTS = ["Feasibility\nStudy","Literature\nSurvey","Interface\nDesign",
            "Database\nDevelopment","Module\nIntegration",
            "Testing and\nDeployment","Documentation"]
    for ai, act in enumerate(ACTS):
        tr,br = ar(ai,0), ar(ai,1)
        ymid = (ROW_Y[tr]+ROW_Y[br+1])/2
        ax.text(COL_X[0]+0.15, ymid, act, ha='left', va='center',
                fontsize=10.5, fontweight='bold', color=BLACK, linespacing=1.3, zorder=3)

    # Legend
    for lx, lc, lt in [(3.5, R, "Planned Schedule"), (6.2, G, "Actual Progress")]:
        rect(lx, TH+0.1, 0.28, 0.22, lc, z=3)
        ax.text(lx+0.38, TH+0.21, lt, va='center', fontsize=9.5, color=BLACK, zorder=3)

    ax.set_ylim(0, TH+0.45)
    plt.tight_layout(pad=0)
    plt.savefig('diagrams/gantt_chart.png', dpi=180, bbox_inches='tight',
                facecolor=WHITE, edgecolor='none')
    plt.close()
    print("  ✓ gantt_chart.png")


# ═══════════════════════════════════════════════════════════════
# 2. ARCHITECTURE DIAGRAM (System Data Flow)
# ═══════════════════════════════════════════════════════════════
def make_architecture():
    steps = [
        ("Camera\nInput",       "Webcam /\ncv2.VideoCapture"),
        ("Pre-\nprocessing",    "Resize &\nNormalize"),
        ("Face\nDetection",     "YuNet CNN\n(ONNX)"),
        ("Face\nRecognition",   "SFace CNN\n(Cosine Sim)"),
        ("Gait\nDetection",     "MOG2\nBackground Sub."),
        ("Output\n& Log",       "GUI + CSV\nLog"),
    ]
    n = len(steps)
    BW, BH = 1.55, 1.15
    GAP = 0.38
    total_w = n*BW + (n-1)*GAP + 1.0
    fig, ax = plt.subplots(figsize=(total_w, 3.2))
    ax.set_xlim(0, total_w); ax.set_ylim(0, 3.2)
    ax.axis('off'); fig.patch.set_facecolor(WHITE)

    for i, (top, bot) in enumerate(steps):
        x = 0.5 + i*(BW+GAP)
        y = 1.0
        # box shadow
        ax.add_patch(FancyBboxPatch((x+0.06, y-0.06), BW, BH,
                     boxstyle="round,pad=0.05",
                     facecolor='#C0C0C0', edgecolor='none', zorder=1))
        # main box
        ax.add_patch(FancyBboxPatch((x, y), BW, BH,
                     boxstyle="round,pad=0.05",
                     facecolor=LGRAY, edgecolor=DBLU, linewidth=2, zorder=2))
        # blue top stripe
        ax.add_patch(FancyBboxPatch((x, y+BH-0.22), BW, 0.22,
                     boxstyle="round,pad=0.0",
                     facecolor=MBLU, edgecolor='none', zorder=3))
        ax.text(x+BW/2, y+BH+0.04, top,
                ha='center', va='bottom', fontsize=10.5,
                fontweight='bold', color=DBLU, linespacing=1.2, zorder=4)
        ax.text(x+BW/2, y+BH/2-0.05, bot,
                ha='center', va='center', fontsize=9,
                color=GRAY, linespacing=1.3, zorder=4)
        if i < n-1:
            ax.annotate('', xy=(x+BW+GAP, y+BH/2), xytext=(x+BW, y+BH/2),
                        arrowprops=dict(arrowstyle='->', color=DBLU,
                                        lw=2.2, mutation_scale=18), zorder=5)

    ax.text(total_w/2, 0.35,
            "Fig 2: System Architecture — Data Flow",
            ha='center', fontsize=10, color=GRAY, style='italic')

    plt.tight_layout(pad=0.3)
    plt.savefig('diagrams/architecture_diagram.png', dpi=180, bbox_inches='tight',
                facecolor=WHITE, edgecolor='none')
    plt.close()
    print("  ✓ architecture_diagram.png")


# ═══════════════════════════════════════════════════════════════
# 3. DATA PROCESSING WORKFLOW
# ═══════════════════════════════════════════════════════════════
def make_data_workflow():
    steps = [
        ("Data\nCollection",      "Webcam / Dataset\nImages"),
        ("Pre-\nprocessing",      "Resize 112×112\nNormalize BGR"),
        ("Feature\nExtraction",   "CNN Embedding\n(YuNet + SFace)"),
        ("Feature\nRepresent.",   "128-dim\nFloat Vector"),
        ("Classifier",            "Cosine Similarity\n≥ 0.363 → Match"),
        ("Output",                "Identity + Gait\nLabel + CSV Log"),
    ]
    COLORS = [MBLU, '#2196F3', '#4CAF50', '#FF9800', '#9C27B0', DBLU]
    n = len(steps)
    BW, BH = 1.55, 1.15
    GAP = 0.38
    total_w = n*BW + (n-1)*GAP + 1.0
    fig, ax = plt.subplots(figsize=(total_w, 3.2))
    ax.set_xlim(0, total_w); ax.set_ylim(0, 3.2)
    ax.axis('off'); fig.patch.set_facecolor(WHITE)

    for i, ((top, bot), col) in enumerate(zip(steps, COLORS)):
        x = 0.5 + i*(BW+GAP)
        y = 1.0
        ax.add_patch(FancyBboxPatch((x+0.06, y-0.06), BW, BH,
                     boxstyle="round,pad=0.05",
                     facecolor='#C0C0C0', edgecolor='none', zorder=1))
        ax.add_patch(FancyBboxPatch((x, y), BW, BH,
                     boxstyle="round,pad=0.05",
                     facecolor=WHITE, edgecolor=col, linewidth=2.5, zorder=2))
        ax.add_patch(FancyBboxPatch((x, y+BH-0.22), BW, 0.22,
                     boxstyle="round,pad=0.0",
                     facecolor=col, edgecolor='none', zorder=3))
        ax.text(x+BW/2, y+BH+0.04, top,
                ha='center', va='bottom', fontsize=10.5,
                fontweight='bold', color=col, linespacing=1.2, zorder=4)
        ax.text(x+BW/2, y+BH/2-0.05, bot,
                ha='center', va='center', fontsize=9,
                color='#404040', linespacing=1.3, zorder=4)
        if i < n-1:
            ax.annotate('', xy=(x+BW+GAP, y+BH/2), xytext=(x+BW, y+BH/2),
                        arrowprops=dict(arrowstyle='->', color='#404040',
                                        lw=2, mutation_scale=18), zorder=5)

    ax.text(total_w/2, 0.35,
            "Fig 3: Data Processing Pipeline",
            ha='center', fontsize=10, color=GRAY, style='italic')

    plt.tight_layout(pad=0.3)
    plt.savefig('diagrams/data_processing_workflow.png', dpi=180, bbox_inches='tight',
                facecolor=WHITE, edgecolor='none')
    plt.close()
    print("  ✓ data_processing_workflow.png")


# ═══════════════════════════════════════════════════════════════
# 4. TEAM STRUCTURE
# ═══════════════════════════════════════════════════════════════
def make_team_structure():
    fig, ax = plt.subplots(figsize=(10, 5.5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 5.5)
    ax.invert_yaxis(); ax.axis('off')
    fig.patch.set_facecolor(WHITE)

    def card(cx, cy, lines, width=2.6, height=1.4, bg=LGRAY, border=DBLU):
        x, y = cx - width/2, cy - height/2
        ax.add_patch(FancyBboxPatch((x, y), width, height,
                     boxstyle="round,pad=0.1",
                     facecolor=bg, edgecolor=border, linewidth=2, zorder=2))
        ax.add_patch(FancyBboxPatch((x, y), width, 0.28,
                     boxstyle="round,pad=0.0",
                     facecolor=border, edgecolor='none', zorder=3))
        for li, (txt, fsz, fw, fc) in enumerate(lines):
            off = 0.05 + li * 0.32
            ax.text(cx, y + 0.38 + off, txt,
                    ha='center', va='top', fontsize=fsz,
                    fontweight=fw, color=fc, zorder=4)

    # Internal Guide
    card(2.2, 1.6,
         [("Internal Guide",      9,  'normal', GRAY),
          ("Mr. Simanta Kalita",  11, 'bold',   DBLU),
          ("Asst. Professor, SMIT", 9, 'normal', BLACK)])

    # External Guide
    card(7.8, 1.6,
         [("External Guide",     9,  'normal', GRAY),
          ("Mr. Sachin Manger",  11, 'bold',   DBLU),
          ("Senior Faculty",      9,  'normal', BLACK)])

    # Student
    card(5.0, 4.0,
         [("Student",            9,  'normal', GRAY),
          ("Afsana Chettri",     11, 'bold',   DBLU),
          ("(202422018), MCA",    9,  'normal', BLACK)])

    # Connector lines
    lw = 2.0
    lc = GRAY
    # Each guide down to join point at y=2.7
    ax.plot([2.2, 2.2], [2.3, 2.9], '-', color=lc, lw=lw, zorder=1)
    ax.plot([7.8, 7.8], [2.3, 2.9], '-', color=lc, lw=lw, zorder=1)
    # Horizontal join
    ax.plot([2.2, 7.8], [2.9, 2.9], '-', color=lc, lw=lw, zorder=1)
    # Down to student
    ax.plot([5.0, 5.0], [2.9, 3.3], '-', color=lc, lw=lw, zorder=1)
    ax.annotate('', xy=(5.0, 3.3), xytext=(5.0, 2.85),
                arrowprops=dict(arrowstyle='->', color=lc, lw=lw,
                                mutation_scale=16), zorder=5)

    ax.text(5.0, 5.25, "Fig 1: Project Team Structure",
            ha='center', fontsize=10, color=GRAY, style='italic')

    plt.tight_layout(pad=0.3)
    plt.savefig('diagrams/team_structure.png', dpi=180, bbox_inches='tight',
                facecolor=WHITE, edgecolor='none')
    plt.close()
    print("  ✓ team_structure.png")


# ═══════════════════════════════════════════════════════════════
# GENERATE ALL PNGS
# ═══════════════════════════════════════════════════════════════
print("Generating diagrams...")
make_gantt()
make_architecture()
make_data_workflow()
make_team_structure()


# ═══════════════════════════════════════════════════════════════
# BUILD WORD DOCUMENT
# ═══════════════════════════════════════════════════════════════
print("\nBuilding Word document...")

doc = Document()

# Page setup — landscape A4-ish
section = doc.sections[0]
section.page_width  = Inches(11.69)
section.page_height = Inches(8.27)
section.left_margin   = Inches(0.8)
section.right_margin  = Inches(0.8)
section.top_margin    = Inches(0.7)
section.bottom_margin = Inches(0.7)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.color.rgb = DRGBColor(0x1F, 0x49, 0x7D)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    return p

def insert_img(path, width_in=9.5, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_in))
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.runs[0]
        cr.font.italic = True
        cr.font.size = Pt(10)
        cr.font.color.rgb = DRGBColor(0x60, 0x60, 0x60)
    doc.add_paragraph()   # spacer

# ── Title ──────────────────────────────────────────────────────
doc.add_paragraph()
heading("FACEE — Deep Learning Approaches for Face & Gait Recognition", level=1)
heading("MCA Major Project — Project Diagrams", level=2)
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Afsana Chettri (202422018)   |   Guide: Mr. Simanta Kalita   |   Co-Guide: Mr. Sachin Manger")
r.font.size = Pt(11); r.font.bold = True
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Dept. of Computer Application, SMIT, Majhitar, East Sikkim")
r2.font.size = Pt(10); r2.font.italic = True; r2.font.color.rgb = DRGBColor(0x60,0x60,0x60)
doc.add_page_break()

# ── Diagram 1: Team Structure ──────────────────────────────────
heading("Fig 1 — Project Team Structure", level=1)
body("The following diagram shows the team hierarchy for the FACEE project, "
     "including the student, internal guide, and external guide.")
insert_img('diagrams/team_structure.png', width_in=7.5,
           caption="Fig 1: Project Team Structure — Afsana Chettri, Mr. Simanta Kalita, Mr. Sachin Manger")
doc.add_page_break()

# ── Diagram 2: Gantt Chart ─────────────────────────────────────
heading("Fig 2 — Project Gantt Chart", level=1)
body("The Gantt chart shows the planned schedule (red rows) and actual progress "
     "(green rows) for each project phase across the timeline from January to July.")
insert_img('diagrams/gantt_chart.png', width_in=9.5,
           caption="Fig 2: Project Gantt Chart — Red = Planned Schedule, Green = Actual Progress")
doc.add_page_break()

# ── Diagram 3: Architecture ────────────────────────────────────
heading("Fig 3 — System Architecture Diagram", level=1)
body("The architecture diagram illustrates the end-to-end data flow of the FACEE "
     "system — from camera input through preprocessing, CNN-based face detection "
     "(YuNet), face recognition (SFace), gait detection (MOG2), and final output logging.")
insert_img('diagrams/architecture_diagram.png', width_in=9.5,
           caption="Fig 3: System Architecture — Data Flow from Camera to Output")
doc.add_page_break()

# ── Diagram 4: Data Processing Workflow ───────────────────────
heading("Fig 4 — Data Processing Workflow", level=1)
body("The data processing workflow diagram shows the six-stage pipeline used to "
     "process input frames: data collection, preprocessing, CNN feature extraction, "
     "128-dimensional embedding representation, cosine similarity classification, "
     "and final identity + gait output.")
insert_img('diagrams/data_processing_workflow.png', width_in=9.5,
           caption="Fig 4: Data Processing Pipeline — From Raw Input to Identified Output")

out = '/home/user/FACEE/FACEE_Diagrams.docx'
doc.save(out)
print(f"\nDone — Word document saved → {out}")
print("Diagrams saved in diagrams/ folder.")
